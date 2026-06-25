"""
Data Dictionary Generator - Tutor School System
สร้างไฟล์ Data Dictionary เป็น Word Document (.docx)
ฟอนต์: TH SarabunPSK | ภาษาไทย

ติดตั้ง: pip install python-docx
รัน:     python generate_data_dictionary.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FONT_NAME = "TH SarabunPSK"
HEADER_BG = "1F4E79"   # สีน้ำเงินเข้ม (header row)
ALT_BG    = "D6E4F0"   # สีฟ้าอ่อน (แถวคู่)
WHITE     = "FFFFFF"
TITLE_COLOR = "1F4E79"

# ============================================================
# Helper Functions
# ============================================================

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="AAAAAA"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, color=None, size=14, align=WD_ALIGN_PARAGRAPH.LEFT):
    for para in cell.paragraphs:
        cell._tc.remove(para._p)
    para = cell.add_paragraph()
    para.alignment = align
    run = para.add_run(str(text))
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # ตั้ง language เป็น Thai
    rPr = run._r.get_or_add_rPr()
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "th-TH")
    rPr.append(lang)
    return para

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

# Column widths (cm) — landscape A4 content width ~24.7 cm
COL_WIDTHS = [1.2, 4.5, 3.0, 1.8, 10.2, 4.0]
COL_HEADERS = ["NO", "Attribute", "Type", "Size", "คำอธิบาย", "Key"]

def add_table(doc, table_name, thai_name, rows, fk_notes=None):
    """สร้างตาราง Data Dictionary หนึ่งตาราง"""

    # ชื่อตาราง
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.add_run(f"ตารางที่ {thai_name}  ({table_name})")
    heading_run.font.name = FONT_NAME
    heading_run.font.size = Pt(18)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor.from_string(TITLE_COLOR)

    # สร้างตาราง
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, (hdr_text, w) in enumerate(zip(COL_HEADERS, COL_WIDTHS)):
        cell = hdr.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, HEADER_BG)
        set_cell_borders(cell, "FFFFFF")
        cell_text(cell, hdr_text, bold=True, color=WHITE, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        bg = ALT_BG if idx % 2 == 1 else WHITE
        for i, (val, w) in enumerate(zip(row_data, COL_WIDTHS)):
            cell = row.cells[i]
            cell.width = Cm(w)
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "AAAAAA")
            align = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3] else WD_ALIGN_PARAGRAPH.LEFT
            # Key column — ถ้ามี FK ทำตัวเอียง
            is_fk = "FK" in str(val)
            cell_text(cell, val, bold=(i == 0 or "PK" in str(val)),
                      color=("8B0000" if is_fk else ("1F4E79" if "PK" in str(val) else None)),
                      size=13, align=align)

    # FK Notes
    if fk_notes:
        note_para = doc.add_paragraph()
        note_run = note_para.add_run("อธิบาย Foreign Key:")
        note_run.font.name = FONT_NAME
        note_run.font.size = Pt(14)
        note_run.font.bold = True
        note_run.font.color.rgb = RGBColor.from_string("8B0000")
        for note in fk_notes:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(note)
            run.font.name = FONT_NAME
            run.font.size = Pt(13)

    doc.add_paragraph()  # เว้นบรรทัด


# ============================================================
# Data Dictionary Data
# ============================================================

TABLES = []

# ──────────────────────────────────────────────
# 1. users
# ──────────────────────────────────────────────
TABLES.append({
    "name": "users",
    "thai": "1. ผู้ใช้งานระบบ",
    "rows": [
        [1, "id",         "BIGINT",    "8",   "รหัสผู้ใช้งาน (Auto Increment)",                     "PK"],
        [2, "username",   "VARCHAR",   "100", "ชื่อผู้ใช้งานในระบบ",                               "UNIQUE"],
        [3, "email",      "VARCHAR",   "255", "อีเมลผู้ใช้งาน (ใช้เป็น principal สำหรับ JWT)",    "UNIQUE"],
        [4, "password",   "VARCHAR",   "-",   "รหัสผ่านที่เข้ารหัสด้วย BCrypt",                    ""],
        [5, "role",       "VARCHAR",   "20",  "บทบาทผู้ใช้งาน: ADMIN, STUDENT, TUTOR",              ""],
        [6, "is_enabled", "BOOLEAN",   "-",   "สถานะการเปิดใช้งานบัญชี (ค่าเริ่มต้น: true)",      ""],
        [7, "created_at", "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                  ""],
        [8, "updated_at", "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",         ""],
    ],
    "fk_notes": None
})

# ──────────────────────────────────────────────
# 2. students
# ──────────────────────────────────────────────
TABLES.append({
    "name": "students",
    "thai": "2. นักเรียน",
    "rows": [
        [1,  "id",                    "BIGINT",    "8",   "รหัสนักเรียน (Auto Increment)",                        "PK"],
        [2,  "user_id",               "BIGINT",    "8",   "อ้างอิงบัญชีผู้ใช้งาน",                              "FK → users.id"],
        [3,  "student_code",          "VARCHAR",   "50",  "รหัสประจำตัวนักเรียน",                                "UNIQUE"],
        [4,  "first_name",            "VARCHAR",   "100", "ชื่อจริง",                                             ""],
        [5,  "last_name",             "VARCHAR",   "100", "นามสกุล",                                              ""],
        [6,  "full_name",             "VARCHAR",   "255", "ชื่อ-นามสกุลเต็ม",                                    ""],
        [7,  "national_id",           "VARCHAR",   "13",  "หมายเลขบัตรประชาชน 13 หลัก",                         "UNIQUE"],
        [8,  "address",               "TEXT",      "-",   "ที่อยู่",                                              ""],
        [9,  "phone_number",          "VARCHAR",   "20",  "เบอร์โทรศัพท์",                                       ""],
        [10, "birth_date",            "DATE",      "-",   "วันเกิด",                                              ""],
        [11, "guardian_phone_number", "VARCHAR",   "20",  "เบอร์โทรผู้ปกครอง",                                   ""],
        [12, "bank_name",             "VARCHAR",   "100", "ชื่อธนาคาร",                                           ""],
        [13, "bank_qr_code",          "TEXT",      "-",   "QR Code พร้อมเพย์ (Base64 หรือ URL)",                 ""],
        [14, "bank_account_name",     "VARCHAR",   "255", "ชื่อบัญชีธนาคาร",                                    ""],
        [15, "bank_account_number",   "VARCHAR",   "50",  "เลขที่บัญชีธนาคาร",                                  ""],
        [16, "created_at",            "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                    ""],
        [17, "updated_at",            "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",           ""],
    ],
    "fk_notes": [
        "user_id → users.id : นักเรียนแต่ละคนผูกกับบัญชีผู้ใช้งาน (users) แบบ One-to-One เพื่อใช้สำหรับการยืนยันตัวตนเข้าสู่ระบบ"
    ]
})

# ──────────────────────────────────────────────
# 3. tutors
# ──────────────────────────────────────────────
TABLES.append({
    "name": "tutors",
    "thai": "3. ครูผู้สอน",
    "rows": [
        [1, "id",             "BIGINT",    "8",   "รหัสครูผู้สอน (Auto Increment)",                    "PK"],
        [2, "user_id",        "BIGINT",    "8",   "อ้างอิงบัญชีผู้ใช้งาน",                           "FK → users.id"],
        [3, "first_name",     "VARCHAR",   "100", "ชื่อจริง",                                          ""],
        [4, "last_name",      "VARCHAR",   "100", "นามสกุล",                                           ""],
        [5, "phone_number",   "VARCHAR",   "20",  "เบอร์โทรศัพท์",                                    ""],
        [6, "specialization", "VARCHAR",   "200", "ความเชี่ยวชาญหรือสาขาวิชาที่สอน",                 ""],
        [7, "bio",            "TEXT",      "-",   "ประวัติย่อของครูผู้สอน",                           ""],
        [8, "created_at",     "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                 ""],
        [9, "updated_at",     "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",        ""],
    ],
    "fk_notes": [
        "user_id → users.id : ครูผู้สอนแต่ละคนผูกกับบัญชีผู้ใช้งาน (users) แบบ One-to-One เพื่อใช้สำหรับการยืนยันตัวตนเข้าสู่ระบบ"
    ]
})

# ──────────────────────────────────────────────
# 4. courses
# ──────────────────────────────────────────────
TABLES.append({
    "name": "courses",
    "thai": "4. คอร์สเรียน",
    "rows": [
        [1,  "id",                      "BIGINT",        "8",   "รหัสคอร์ส (Auto Increment)",                         "PK"],
        [2,  "course_code",             "VARCHAR",       "50",  "รหัสคอร์ส (ไม่ซ้ำกัน)",                             "UNIQUE"],
        [3,  "course_name",             "VARCHAR",       "200", "ชื่อคอร์ส",                                           ""],
        [4,  "price",                   "DECIMAL(10,2)", "-",   "ราคาคอร์ส (หน่วย: บาท)",                             ""],
        [5,  "description",             "TEXT",          "-",   "รายละเอียดคอร์ส",                                    ""],
        [6,  "total_hours",             "INT",           "4",   "จำนวนชั่วโมงเรียนทั้งหมด",                           ""],
        [7,  "seat_limit",              "INT",           "4",   "จำนวนที่นั่งสูงสุด",                                  ""],
        [8,  "registration_start_date", "DATE",          "-",   "วันเริ่มเปิดรับสมัคร",                               ""],
        [9,  "registration_end_date",   "DATE",          "-",   "วันปิดรับสมัคร",                                     ""],
        [10, "course_start_date",       "DATE",          "-",   "วันเริ่มเรียน",                                       ""],
        [11, "status",                  "VARCHAR",       "30",  "สถานะคอร์ส: DRAFT, PUBLISHED, CLOSED, CANCELLED",   ""],
        [12, "tutor_id",                "BIGINT",        "8",   "อ้างอิงครูผู้สอน",                                   "FK → tutors.id"],
        [13, "tutor_remark",            "TEXT",          "-",   "หมายเหตุจากครูผู้สอน",                               ""],
        [14, "schedule_days",           "VARCHAR",       "100", "วันที่สอนในสัปดาห์ (เช่น MON,WED,FRI)",             ""],
        [15, "schedule_start_time",     "TIME",          "-",   "เวลาเริ่มเรียนประจำสัปดาห์",                         ""],
        [16, "schedule_end_time",       "TIME",          "-",   "เวลาสิ้นสุดการเรียนประจำสัปดาห์",                   ""],
        [17, "created_at",              "TIMESTAMP",     "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                   ""],
        [18, "updated_at",              "TIMESTAMP",     "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",          ""],
    ],
    "fk_notes": [
        "tutor_id → tutors.id : ระบุครูผู้สอนที่รับผิดชอบคอร์สนี้ (Many-to-One) ครูหนึ่งคนสามารถมีหลายคอร์สได้"
    ]
})

# ──────────────────────────────────────────────
# 5. course_lessons
# ──────────────────────────────────────────────
TABLES.append({
    "name": "course_lessons",
    "thai": "5. บทเรียนในคอร์ส",
    "rows": [
        [1, "id",             "BIGINT",    "8",   "รหัสบทเรียน (Auto Increment)",                     "PK"],
        [2, "course_id",      "BIGINT",    "8",   "อ้างอิงคอร์สที่บทเรียนนี้อยู่",                  "FK → courses.id"],
        [3, "lesson_title",   "VARCHAR",   "300", "ชื่อหัวข้อบทเรียน",                               ""],
        [4, "lesson_content", "TEXT",      "-",   "เนื้อหาบทเรียน",                                  ""],
        [5, "lesson_order",   "INT",       "4",   "ลำดับบทเรียนภายในคอร์ส",                         ""],
        [6, "created_at",     "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                ""],
        [7, "updated_at",     "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",       ""],
    ],
    "fk_notes": [
        "course_id → courses.id : บทเรียนสังกัดอยู่ในคอร์ส (Many-to-One) คอร์สหนึ่งมีได้หลายบทเรียน เมื่อลบคอร์สจะลบบทเรียนทั้งหมดด้วย (Cascade)"
    ]
})

# ──────────────────────────────────────────────
# 6. course_tests
# ──────────────────────────────────────────────
TABLES.append({
    "name": "course_tests",
    "thai": "6. แบบทดสอบประจำบทเรียน",
    "rows": [
        [1, "id",               "BIGINT",    "8",   "รหัสแบบทดสอบ (Auto Increment)",                     "PK"],
        [2, "course_id",        "BIGINT",    "8",   "อ้างอิงคอร์สที่แบบทดสอบนี้อยู่",                 "FK → courses.id"],
        [3, "test_title",       "VARCHAR",   "300", "ชื่อแบบทดสอบ",                                     ""],
        [4, "test_description", "TEXT",      "-",   "รายละเอียดแบบทดสอบ",                              ""],
        [5, "test_order",       "INT",       "4",   "ลำดับแบบทดสอบในคอร์ส",                            ""],
        [6, "lesson_order",     "INT",       "4",   "ลำดับบทเรียนที่สัมพันธ์กับแบบทดสอบนี้",          ""],
        [7, "created_at",       "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                 ""],
        [8, "updated_at",       "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",        ""],
    ],
    "fk_notes": [
        "course_id → courses.id : แบบทดสอบสังกัดอยู่ในคอร์ส (Many-to-One) คอร์สหนึ่งมีได้หลายแบบทดสอบ เมื่อลบคอร์สจะลบแบบทดสอบทั้งหมดด้วย (Cascade)"
    ]
})

# ──────────────────────────────────────────────
# 7. test_questions
# ──────────────────────────────────────────────
TABLES.append({
    "name": "test_questions",
    "thai": "7. คำถามในแบบทดสอบ",
    "rows": [
        [1, "id",             "BIGINT",    "8",   "รหัสคำถาม (Auto Increment)",                                               "PK"],
        [2, "course_test_id", "BIGINT",    "8",   "อ้างอิงแบบทดสอบที่คำถามนี้อยู่",                                         "FK → course_tests.id"],
        [3, "question_text",  "TEXT",      "-",   "ข้อความของคำถาม",                                                         ""],
        [4, "question_type",  "VARCHAR",   "30",  "ประเภทคำถาม: MULTIPLE_CHOICE, SHORT_ANSWER, PARAGRAPH, CHECKBOX",         ""],
        [5, "question_order", "INT",       "4",   "ลำดับข้อคำถาม",                                                           ""],
        [6, "explanation",    "TEXT",      "-",   "คำอธิบายเฉลย",                                                            ""],
        [7, "created_at",     "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                        ""],
        [8, "updated_at",     "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                               ""],
    ],
    "fk_notes": [
        "course_test_id → course_tests.id : คำถามสังกัดอยู่ในแบบทดสอบ (Many-to-One) แบบทดสอบหนึ่งมีได้หลายคำถาม เมื่อลบแบบทดสอบจะลบคำถามทั้งหมดด้วย (Cascade)"
    ]
})

# ──────────────────────────────────────────────
# 8. test_question_options
# ──────────────────────────────────────────────
TABLES.append({
    "name": "test_question_options",
    "thai": "8. ตัวเลือกคำถามแบบทดสอบ",
    "rows": [
        [1, "id",           "BIGINT",    "8",   "รหัสตัวเลือก (Auto Increment)",                  "PK"],
        [2, "question_id",  "BIGINT",    "8",   "อ้างอิงคำถามที่ตัวเลือกนี้อยู่",              "FK → test_questions.id"],
        [3, "option_text",  "TEXT",      "-",   "ข้อความของตัวเลือก",                            ""],
        [4, "is_correct",   "BOOLEAN",   "-",   "ระบุว่าตัวเลือกนี้เป็นคำตอบที่ถูกต้องหรือไม่", ""],
        [5, "option_order", "INT",       "4",   "ลำดับตัวเลือก",                                 ""],
        [6, "created_at",   "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",              ""],
        [7, "updated_at",   "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",     ""],
    ],
    "fk_notes": [
        "question_id → test_questions.id : ตัวเลือกสังกัดอยู่ในคำถาม (Many-to-One) คำถามหนึ่งมีได้หลายตัวเลือก เมื่อลบคำถามจะลบตัวเลือกทั้งหมดด้วย (Cascade)"
    ]
})

# ──────────────────────────────────────────────
# 9. course_schedules
# ──────────────────────────────────────────────
TABLES.append({
    "name": "course_schedules",
    "thai": "9. ตารางการสอน",
    "rows": [
        [1,  "id",             "BIGINT",    "8",   "รหัสตารางสอน (Auto Increment)",                                            "PK"],
        [2,  "schedule_code",  "VARCHAR",   "30",  "รหัสตารางสอน (ไม่ซ้ำกัน)",                                               "UNIQUE"],
        [3,  "course_id",      "BIGINT",    "8",   "อ้างอิงคอร์สที่เปิดสอน",                                                 "FK → courses.id"],
        [4,  "lesson_id",      "BIGINT",    "8",   "อ้างอิงบทเรียนที่สอน (อาจเป็น null)",                                   "FK → course_lessons.id"],
        [5,  "tutor_id",       "BIGINT",    "8",   "อ้างอิงครูผู้สอนในรายการนี้",                                            "FK → tutors.id"],
        [6,  "title",          "VARCHAR",   "300", "ชื่อหัวข้อการสอน",                                                        ""],
        [7,  "description",    "TEXT",      "-",   "รายละเอียดการสอน",                                                        ""],
        [8,  "schedule_date",  "DATE",      "-",   "วันที่สอน (ห้ามแก้ไขหลังสร้าง ต้อง Cancel แทน)",                        ""],
        [9,  "start_time",     "TIME",      "-",   "เวลาเริ่มสอน (ห้ามแก้ไขหลังสร้าง)",                                     ""],
        [10, "end_time",       "TIME",      "-",   "เวลาสิ้นสุดการสอน (ห้ามแก้ไขหลังสร้าง)",                                ""],
        [11, "location",       "VARCHAR",   "500", "สถานที่สอน",                                                              ""],
        [12, "meeting_link",   "VARCHAR",   "500", "ลิงก์ประชุมออนไลน์",                                                     ""],
        [13, "schedule_type",  "VARCHAR",   "20",  "ประเภทการสอน: ONLINE, ONSITE, HYBRID",                                   ""],
        [14, "status",         "VARCHAR",   "20",  "สถานะ: SCHEDULED, COMPLETED, CANCELLED",                                  ""],
        [15, "cancel_reason",  "TEXT",      "-",   "เหตุผลในการยกเลิกตารางสอน",                                              ""],
        [16, "cancelled_at",   "TIMESTAMP", "-",   "วันเวลาที่ยกเลิกตารางสอน",                                               ""],
        [17, "cancelled_by",   "VARCHAR",   "100", "ผู้ที่ดำเนินการยกเลิกตารางสอน",                                          ""],
        [18, "created_at",     "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                        ""],
        [19, "updated_at",     "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                               ""],
    ],
    "fk_notes": [
        "course_id → courses.id : ตารางสอนสังกัดอยู่ในคอร์ส (Many-to-One) คอร์สหนึ่งมีได้หลายรายการตารางสอน",
        "lesson_id → course_lessons.id : ระบุบทเรียนที่สอนในรายการนี้ (Many-to-One, Nullable) หากไม่ระบุหมายถึงสอนแบบทั่วไปไม่ผูกกับบทเรียนเฉพาะ",
        "tutor_id → tutors.id : ระบุครูผู้สอนที่รับผิดชอบรายการสอนนี้ (Many-to-One)"
    ]
})

# ──────────────────────────────────────────────
# 10. enrollments
# ──────────────────────────────────────────────
TABLES.append({
    "name": "enrollments",
    "thai": "10. การลงทะเบียนเรียน",
    "rows": [
        [1,  "id",               "BIGINT",        "8",   "รหัสการลงทะเบียน (Auto Increment)",                                        "PK"],
        [2,  "enrollment_code",  "VARCHAR",        "20",  "รหัสการลงทะเบียน (ไม่ซ้ำกัน)",                                            "UNIQUE"],
        [3,  "student_id",       "BIGINT",         "8",   "อ้างอิงนักเรียนที่ลงทะเบียน",                                            "FK → students.id"],
        [4,  "course_id",        "BIGINT",         "8",   "อ้างอิงคอร์สที่ลงทะเบียน",                                               "FK → courses.id"],
        [5,  "enrollment_date",  "TIMESTAMP",      "-",   "วันเวลาที่ลงทะเบียน (Auto Set)",                                          ""],
        [6,  "status",           "VARCHAR",        "20",  "สถานะ: PENDING, APPROVED, REJECTED, CANCELLED",                           ""],
        [7,  "payment_status",   "VARCHAR",        "30",  "สถานะการชำระ: UNPAID, PAID, REFUNDED",                                    ""],
        [8,  "payment_method",   "VARCHAR",        "20",  "วิธีชำระ: CASH, TRANSFER, QR_CODE",                                       ""],
        [9,  "amount",           "DECIMAL(10,2)",  "-",   "ราคาเต็มก่อนส่วนลด (บาท)",                                               ""],
        [10, "discount_amount",  "DECIMAL(10,2)",  "-",   "จำนวนส่วนลด (บาท) ค่าเริ่มต้น 0",                                        ""],
        [11, "final_amount",     "DECIMAL(10,2)",  "-",   "ราคาสุทธิหลังหักส่วนลด (บาท)",                                           ""],
        [12, "payment_slip_url", "VARCHAR",        "500", "URL รูปสลิปการชำระเงิน",                                                 ""],
        [13, "note",             "TEXT",           "-",   "หมายเหตุเพิ่มเติม",                                                       ""],
        [14, "approved_by",      "VARCHAR",        "100", "ผู้อนุมัติการลงทะเบียน",                                                  ""],
        [15, "approved_at",      "TIMESTAMP",      "-",   "วันเวลาที่อนุมัติการลงทะเบียน",                                          ""],
        [16, "created_at",       "TIMESTAMP",      "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                        ""],
        [17, "updated_at",       "TIMESTAMP",      "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                               ""],
    ],
    "fk_notes": [
        "student_id → students.id : ระบุนักเรียนที่ทำการลงทะเบียน (Many-to-One) นักเรียนหนึ่งคนลงทะเบียนได้หลายคอร์ส",
        "course_id → courses.id : ระบุคอร์สที่ถูกลงทะเบียน (Many-to-One) นักเรียนคนเดิมลงทะเบียนคอร์สเดิมซ้ำไม่ได้ (Unique Constraint: student_id + course_id)"
    ]
})

# ──────────────────────────────────────────────
# 11. payments
# ──────────────────────────────────────────────
TABLES.append({
    "name": "payments",
    "thai": "11. การชำระเงิน",
    "rows": [
        [1,  "id",                     "BIGINT",        "8",   "รหัสการชำระเงิน (Auto Increment)",                           "PK"],
        [2,  "payment_code",           "VARCHAR",        "20",  "รหัสการชำระเงิน (ไม่ซ้ำกัน)",                               "UNIQUE"],
        [3,  "enrollment_id",          "BIGINT",         "8",   "อ้างอิงการลงทะเบียนที่ชำระ",                               "FK → enrollments.id"],
        [4,  "student_id",             "BIGINT",         "8",   "อ้างอิงนักเรียนที่ชำระเงิน",                               "FK → students.id"],
        [5,  "institution_profile_id", "BIGINT",         "8",   "อ้างอิงโปรไฟล์สถาบันที่รับชำระ",                          "FK → institution_profiles.id"],
        [6,  "payment_date",           "TIMESTAMP",      "-",   "วันเวลาที่ชำระเงิน (Auto Set)",                             ""],
        [7,  "amount",                 "DECIMAL(10,2)",  "-",   "จำนวนเงินที่ชำระ (บาท)",                                   ""],
        [8,  "payment_method",         "VARCHAR",        "20",  "วิธีการชำระ: CASH, TRANSFER, QR_CODE",                     ""],
        [9,  "payment_status",         "VARCHAR",        "20",  "สถานะตรวจสอบ: PENDING, VERIFIED, REJECTED",               ""],
        [10, "payment_slip_url",       "VARCHAR",        "500", "URL รูปสลิปการชำระเงิน",                                   ""],
        [11, "transaction_reference",  "VARCHAR",        "100", "หมายเลขอ้างอิงธุรกรรม",                                   ""],
        [12, "note",                   "TEXT",           "-",   "หมายเหตุเพิ่มเติม",                                        ""],
        [13, "verified_by",            "VARCHAR",        "100", "ผู้ตรวจสอบการชำระเงิน",                                   ""],
        [14, "verified_at",            "TIMESTAMP",      "-",   "วันเวลาที่ตรวจสอบการชำระเงิน",                             ""],
        [15, "created_at",             "TIMESTAMP",      "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                         ""],
        [16, "updated_at",             "TIMESTAMP",      "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                ""],
    ],
    "fk_notes": [
        "enrollment_id → enrollments.id : การชำระเงินผูกกับรายการลงทะเบียน (Many-to-One) หนึ่งการลงทะเบียนมีได้หลายรายการชำระ (กรณีชำระบางส่วน)",
        "student_id → students.id : ระบุนักเรียนที่ชำระเงิน (Many-to-One) เพื่อให้ query ประวัติชำระของนักเรียนได้โดยตรง",
        "institution_profile_id → institution_profiles.id : ระบุบัญชีธนาคารของสถาบันที่ใช้รับชำระ (Many-to-One) เพื่อให้ระบบทราบข้อมูลบัญชีปลายทาง"
    ]
})

# ──────────────────────────────────────────────
# 12. institution_profiles
# ──────────────────────────────────────────────
TABLES.append({
    "name": "institution_profiles",
    "thai": "12. โปรไฟล์สถาบัน",
    "rows": [
        [1,  "id",                   "BIGINT",    "8",   "รหัสโปรไฟล์สถาบัน (Auto Increment)",               "PK"],
        [2,  "institution_code",     "VARCHAR",   "50",  "รหัสสถาบัน (ไม่ซ้ำกัน)",                          "UNIQUE"],
        [3,  "institution_name",     "VARCHAR",   "255", "ชื่อสถาบัน",                                        ""],
        [4,  "address",              "TEXT",      "-",   "ที่อยู่สถาบัน",                                     ""],
        [5,  "phone_number",         "VARCHAR",   "20",  "เบอร์โทรศัพท์สถาบัน",                              ""],
        [6,  "email",                "VARCHAR",   "255", "อีเมลติดต่อสถาบัน",                                ""],
        [7,  "logo_url",             "TEXT",      "-",   "URL โลโก้สถาบัน",                                  ""],
        [8,  "bank_name",            "VARCHAR",   "100", "ชื่อธนาคารสถาบัน",                                 ""],
        [9,  "bank_account_name",    "VARCHAR",   "255", "ชื่อบัญชีธนาคารสถาบัน",                           ""],
        [10, "bank_account_number",  "VARCHAR",   "50",  "เลขที่บัญชีธนาคารสถาบัน",                         ""],
        [11, "bank_qr_code",         "TEXT",      "-",   "QR Code พร้อมเพย์สถาบัน (Base64 หรือ URL)",       ""],
        [12, "created_at",           "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                 ""],
        [13, "updated_at",           "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",        ""],
    ],
    "fk_notes": None
})

# ──────────────────────────────────────────────
# 13. classroom_sessions
# ──────────────────────────────────────────────
TABLES.append({
    "name": "classroom_sessions",
    "thai": "13. Session ห้องเรียนออนไลน์",
    "rows": [
        [1,  "id",                    "BIGINT",    "8",   "รหัส Session ห้องเรียน (Auto Increment)",                      "PK"],
        [2,  "session_code",          "VARCHAR",   "50",  "รหัส Session (ไม่ซ้ำกัน)",                                    "UNIQUE"],
        [3,  "course_id",             "BIGINT",    "8",   "อ้างอิงคอร์สที่เปิด Session",                                "FK → courses.id"],
        [4,  "lesson_id",             "BIGINT",    "8",   "อ้างอิงบทเรียนที่สอนใน Session (อาจเป็น null)",            "FK → course_lessons.id"],
        [5,  "tutor_id",              "BIGINT",    "8",   "อ้างอิงครูผู้เปิด Session",                                  "FK → tutors.id"],
        [6,  "start_time",            "TIMESTAMP", "-",   "วันเวลาเริ่ม Session",                                        ""],
        [7,  "end_time",              "TIMESTAMP", "-",   "วันเวลาสิ้นสุด Session",                                     ""],
        [8,  "late_threshold_minutes","INT",        "4",   "นาทีที่ถือว่ามาสาย (ค่าเริ่มต้น: 15 นาที)",                ""],
        [9,  "join_code",             "VARCHAR",   "100", "รหัสลับสำหรับเข้าร่วม Session",                              ""],
        [10, "is_camera_required",    "BOOLEAN",   "-",   "บังคับให้เปิดกล้องหรือไม่ (ค่าเริ่มต้น: false)",           ""],
        [11, "status",                "VARCHAR",   "20",  "สถานะ: SCHEDULED, OPEN, CLOSED, CANCELLED",                  ""],
        [12, "created_at",            "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                           ""],
        [13, "updated_at",            "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                  ""],
    ],
    "fk_notes": [
        "course_id → courses.id : Session สังกัดอยู่ในคอร์ส (Many-to-One) คอร์สหนึ่งมีได้หลาย Session",
        "lesson_id → course_lessons.id : ระบุบทเรียนที่สอนใน Session นี้ (Many-to-One, Nullable) หากไม่ระบุหมายถึงสอนแบบรวม",
        "tutor_id → tutors.id : ระบุครูผู้เปิดและดูแล Session (Many-to-One)"
    ]
})

# ──────────────────────────────────────────────
# 14. attendance_records
# ──────────────────────────────────────────────
TABLES.append({
    "name": "attendance_records",
    "thai": "14. บันทึกการเข้าเรียน",
    "rows": [
        [1,  "id",                  "BIGINT",    "8",   "รหัสบันทึกการเข้าเรียน (Auto Increment)",                        "PK"],
        [2,  "attendance_code",     "VARCHAR",   "50",  "รหัสบันทึก (ไม่ซ้ำกัน)",                                        "UNIQUE"],
        [3,  "enrollment_id",       "BIGINT",    "8",   "อ้างอิงการลงทะเบียน",                                           "FK → enrollments.id"],
        [4,  "student_id",          "BIGINT",    "8",   "อ้างอิงนักเรียน",                                               "FK → students.id"],
        [5,  "course_id",           "BIGINT",    "8",   "อ้างอิงคอร์ส",                                                  "FK → courses.id"],
        [6,  "lesson_id",           "BIGINT",    "8",   "อ้างอิงบทเรียน (อาจเป็น null)",                                "FK → course_lessons.id"],
        [7,  "session_id",          "BIGINT",    "8",   "อ้างอิง Session ห้องเรียน",                                    "FK → classroom_sessions.id"],
        [8,  "first_join_at",       "TIMESTAMP", "-",   "วันเวลาที่เข้าร่วมครั้งแรก (Auto Set, ห้ามแก้ไข)",            ""],
        [9,  "last_leave_at",       "TIMESTAMP", "-",   "วันเวลาที่ออกจาก Session ครั้งสุดท้าย",                       ""],
        [10, "check_in_time",       "TIMESTAMP", "-",   "เวลา Check-in อย่างเป็นทางการ (Auto Set, ห้ามแก้ไข)",         ""],
        [11, "check_out_time",      "TIMESTAMP", "-",   "เวลา Check-out อย่างเป็นทางการ",                               ""],
        [12, "late_minutes",        "INT",        "4",   "จำนวนนาทีที่มาสาย (ค่าเริ่มต้น: 0)",                          ""],
        [13, "status",              "VARCHAR",   "20",  "สถานะ: PRESENT, LATE, ABSENT, EXCUSED",                         ""],
        [14, "attendance_method",   "VARCHAR",   "20",  "วิธีบันทึก: AUTO_JOIN, MANUAL",                                 ""],
        [15, "camera_verified",     "BOOLEAN",   "-",   "ผ่านการตรวจสอบกล้องหรือไม่ (ค่าเริ่มต้น: false)",            ""],
        [16, "camera_snapshot_url", "TEXT",      "-",   "URL ภาพถ่ายจากกล้องที่ใช้ยืนยันตัวตน",                       ""],
        [17, "note",                "TEXT",      "-",   "หมายเหตุเพิ่มเติม",                                             ""],
        [18, "created_at",          "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                              ""],
        [19, "updated_at",          "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                    ""],
    ],
    "fk_notes": [
        "enrollment_id → enrollments.id : บันทึกการเข้าเรียนผูกกับการลงทะเบียน (Many-to-One) เพื่อยืนยันว่านักเรียนมีสิทธิ์เข้าเรียน",
        "student_id → students.id : ระบุนักเรียนที่เข้าเรียน (Many-to-One) นักเรียนกับ Session เดียวกันมีได้แค่ 1 บันทึก (Unique: student_id + session_id)",
        "course_id → courses.id : ระบุคอร์สที่เข้าเรียน (Many-to-One) เก็บไว้เพื่อ Query ประวัติเข้าเรียนตามคอร์สได้โดยตรง",
        "lesson_id → course_lessons.id : ระบุบทเรียนที่เข้าเรียน (Many-to-One, Nullable)",
        "session_id → classroom_sessions.id : ระบุ Session ที่บันทึกการเข้าเรียน (Many-to-One)"
    ]
})

# ──────────────────────────────────────────────
# 15. attendance_audit_logs
# ──────────────────────────────────────────────
TABLES.append({
    "name": "attendance_audit_logs",
    "thai": "15. บันทึกการแก้ไขสถานะการเข้าเรียน",
    "rows": [
        [1, "id",                    "BIGINT",    "8",   "รหัส Log (Auto Increment)",                                       "PK"],
        [2, "attendance_record_id",  "BIGINT",    "8",   "อ้างอิงบันทึกการเข้าเรียนที่ถูกแก้ไข",                         "FK → attendance_records.id"],
        [3, "old_status",            "VARCHAR",   "20",  "สถานะเดิมก่อนแก้ไข: PRESENT, LATE, ABSENT, EXCUSED",            ""],
        [4, "new_status",            "VARCHAR",   "20",  "สถานะใหม่หลังแก้ไข: PRESENT, LATE, ABSENT, EXCUSED",            ""],
        [5, "changed_by",            "VARCHAR",   "255", "ผู้ที่ดำเนินการแก้ไขสถานะ",                                     ""],
        [6, "changed_role",          "VARCHAR",   "50",  "บทบาทของผู้แก้ไข เช่น ADMIN, TUTOR",                            ""],
        [7, "reason",                "TEXT",      "-",   "เหตุผลในการแก้ไขสถานะ",                                          ""],
        [8, "changed_at",            "TIMESTAMP", "-",   "วันเวลาที่แก้ไขสถานะ (Auto Set)",                                ""],
    ],
    "fk_notes": [
        "attendance_record_id → attendance_records.id : Log ผูกกับบันทึกการเข้าเรียน (Many-to-One) ทุกครั้งที่สถานะเปลี่ยนจะสร้าง Log ใหม่ เพื่อความโปร่งใส"
    ]
})

# ──────────────────────────────────────────────
# 16. course_evaluations
# ──────────────────────────────────────────────
TABLES.append({
    "name": "course_evaluations",
    "thai": "16. การประเมินคอร์ส",
    "rows": [
        [1,  "id",                   "BIGINT",    "8",   "รหัสการประเมิน (Auto Increment)",                               "PK"],
        [2,  "evaluation_code",      "VARCHAR",   "20",  "รหัสการประเมิน (ไม่ซ้ำกัน)",                                  "UNIQUE"],
        [3,  "student_id",           "BIGINT",    "8",   "อ้างอิงนักเรียนผู้ประเมิน",                                    "FK → students.id"],
        [4,  "course_id",            "BIGINT",    "8",   "อ้างอิงคอร์สที่ถูกประเมิน",                                   "FK → courses.id"],
        [5,  "enrollment_id",        "BIGINT",    "8",   "อ้างอิงการลงทะเบียน (ตรวจสอบสิทธิ์ประเมิน)",                 "FK → enrollments.id"],
        [6,  "tutor_id",             "BIGINT",    "8",   "อ้างอิงครูผู้สอนที่ถูกประเมิน",                              "FK → tutors.id"],
        [7,  "rating",               "INT",       "4",   "คะแนนภาพรวม (1-5 ดาว)",                                        ""],
        [8,  "teaching_score",       "INT",       "4",   "คะแนนด้านการสอน (1-5)",                                        ""],
        [9,  "content_score",        "INT",       "4",   "คะแนนด้านเนื้อหา (1-5)",                                       ""],
        [10, "material_score",       "INT",       "4",   "คะแนนด้านสื่อการสอน (1-5)",                                   ""],
        [11, "communication_score",  "INT",       "4",   "คะแนนด้านการสื่อสาร (1-5)",                                   ""],
        [12, "value_score",          "INT",       "4",   "คะแนนด้านความคุ้มค่า (1-5)",                                  ""],
        [13, "comment",              "TEXT",      "-",   "ความคิดเห็นเพิ่มเติม",                                         ""],
        [14, "suggestion",           "TEXT",      "-",   "ข้อเสนอแนะสำหรับปรับปรุง",                                    ""],
        [15, "is_anonymous",         "BOOLEAN",   "-",   "ประเมินแบบนิรนามหรือไม่ (ค่าเริ่มต้น: false)",               ""],
        [16, "status",               "VARCHAR",   "20",  "สถานะ: PENDING, APPROVED, REJECTED",                           ""],
        [17, "submitted_at",         "TIMESTAMP", "-",   "วันเวลาที่ส่งการประเมิน (Auto Set)",                           ""],
        [18, "created_at",           "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                             ""],
        [19, "updated_at",           "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                   ""],
    ],
    "fk_notes": [
        "student_id → students.id : ระบุนักเรียนผู้ประเมิน (Many-to-One) นักเรียนประเมินคอร์สเดิมซ้ำไม่ได้ (Unique: student_id + course_id)",
        "course_id → courses.id : ระบุคอร์สที่ถูกประเมิน (Many-to-One)",
        "enrollment_id → enrollments.id : ตรวจสอบว่านักเรียนลงทะเบียนและมีสิทธิ์ประเมิน (Many-to-One)",
        "tutor_id → tutors.id : ระบุครูผู้สอนที่ถูกประเมิน (Many-to-One)"
    ]
})

# ──────────────────────────────────────────────
# 17. notifications
# ──────────────────────────────────────────────
TABLES.append({
    "name": "notifications",
    "thai": "17. การแจ้งเตือน",
    "rows": [
        [1,  "id",                "BIGINT",    "8",   "รหัสการแจ้งเตือน (Auto Increment)",                                                       "PK"],
        [2,  "notification_code", "VARCHAR",   "30",  "รหัสการแจ้งเตือน (ไม่ซ้ำกัน)",                                                           "UNIQUE"],
        [3,  "user_id",           "BIGINT",    "8",   "อ้างอิงผู้ใช้งานที่รับการแจ้งเตือน (Nullable)",                                          "FK → users.id"],
        [4,  "recipient_email",   "VARCHAR",   "255", "อีเมลผู้รับการแจ้งเตือน",                                                                 ""],
        [5,  "subject",           "VARCHAR",   "500", "หัวข้อการแจ้งเตือน",                                                                       ""],
        [6,  "message",           "TEXT",      "-",   "เนื้อหาข้อความแจ้งเตือน",                                                                 ""],
        [7,  "notification_type", "VARCHAR",   "40",  "ประเภท: ENROLLMENT, PAYMENT, ATTENDANCE, EXAM ฯลฯ",                                        ""],
        [8,  "reference_type",    "VARCHAR",   "30",  "ประเภทข้อมูลที่อ้างอิง: NONE, ENROLLMENT, PAYMENT, SESSION ฯลฯ",                          ""],
        [9,  "reference_id",      "BIGINT",    "8",   "รหัสของข้อมูลที่อ้างอิง (ไม่มี FK constraint ตรงๆ เก็บเป็น logical reference)",          ""],
        [10, "delivery_channel",  "VARCHAR",   "20",  "ช่องทางการส่ง: EMAIL (ค่าเริ่มต้น)",                                                      ""],
        [11, "delivery_status",   "VARCHAR",   "20",  "สถานะการส่ง: PENDING, SENT, FAILED",                                                       ""],
        [12, "sent_at",           "TIMESTAMP", "-",   "วันเวลาที่ส่งสำเร็จ",                                                                     ""],
        [13, "failed_reason",     "TEXT",      "-",   "เหตุผลที่ส่งไม่สำเร็จ",                                                                   ""],
        [14, "created_at",        "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                                        ""],
        [15, "updated_at",        "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                                               ""],
    ],
    "fk_notes": [
        "user_id → users.id : อ้างอิงผู้ใช้งานที่เป็นเจ้าของการแจ้งเตือน (Many-to-One, Nullable) หากเป็น null หมายความว่าแจ้งเตือนตาม email โดยตรงโดยไม่ผูกกับบัญชีในระบบ",
        "reference_id : เป็น Logical Foreign Key ที่ไม่มี DB Constraint — ใช้ร่วมกับ reference_type เพื่ออ้างอิงข้อมูลในตารางต่างๆ เช่น ENROLLMENT → enrollments.id, PAYMENT → payments.id"
    ]
})

# ──────────────────────────────────────────────
# 18. exams
# ──────────────────────────────────────────────
TABLES.append({
    "name": "exams",
    "thai": "18. ข้อสอบ",
    "rows": [
        [1,  "id",                              "BIGINT",    "8",   "รหัสข้อสอบ (Auto Increment)",                                          "PK"],
        [2,  "exam_code",                       "VARCHAR",   "50",  "รหัสข้อสอบ (ไม่ซ้ำกัน)",                                             "UNIQUE"],
        [3,  "course_id",                       "BIGINT",    "8",   "อ้างอิงคอร์สที่ข้อสอบนี้สังกัด",                                    "FK → courses.id"],
        [4,  "lesson_id",                       "BIGINT",    "8",   "อ้างอิงบทเรียนที่ข้อสอบนี้สังกัด (อาจเป็น null)",                 "FK → course_lessons.id"],
        [5,  "tutor_id",                        "BIGINT",    "8",   "อ้างอิงครูผู้สร้างข้อสอบ",                                          "FK → tutors.id"],
        [6,  "title",                           "VARCHAR",   "300", "ชื่อข้อสอบ",                                                         ""],
        [7,  "description",                     "TEXT",      "-",   "รายละเอียดและคำชี้แจงข้อสอบ",                                       ""],
        [8,  "total_score",                     "DOUBLE",    "-",   "คะแนนรวมทั้งหมด (คำนวณอัตโนมัติ ค่าเริ่มต้น: 0)",                 ""],
        [9,  "passing_score",                   "DOUBLE",    "-",   "คะแนนขั้นต่ำที่ถือว่าผ่าน",                                        ""],
        [10, "start_time",                      "TIMESTAMP", "-",   "วันเวลาที่เปิดให้ทำข้อสอบ",                                        ""],
        [11, "end_time",                        "TIMESTAMP", "-",   "วันเวลาที่ปิดรับการส่งข้อสอบ",                                     ""],
        [12, "duration_minutes",                "INT",       "4",   "เวลาทำข้อสอบสูงสุด (หน่วย: นาที)",                                 ""],
        [13, "allow_multiple_attempts",         "BOOLEAN",   "-",   "อนุญาตให้ทำข้อสอบได้หลายครั้งหรือไม่",                            ""],
        [14, "max_attempts",                    "INT",       "4",   "จำนวนครั้งสูงสุดที่อนุญาต",                                        ""],
        [15, "shuffle_questions",               "BOOLEAN",   "-",   "สุ่มลำดับข้อคำถามหรือไม่",                                         ""],
        [16, "show_score_after_submit",         "BOOLEAN",   "-",   "แสดงคะแนนทันทีหลังส่งหรือไม่ (ค่าเริ่มต้น: true)",               ""],
        [17, "show_correct_answers_after_submit","BOOLEAN",   "-",   "แสดงเฉลยหลังส่งข้อสอบหรือไม่ (ค่าเริ่มต้น: false)",              ""],
        [18, "status",                          "VARCHAR",   "20",  "สถานะข้อสอบ: DRAFT, PUBLISHED, CLOSED",                             ""],
        [19, "created_at",                      "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                  ""],
        [20, "updated_at",                      "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                         ""],
    ],
    "fk_notes": [
        "course_id → courses.id : ข้อสอบสังกัดคอร์ส (Many-to-One) คอร์สหนึ่งมีได้หลายข้อสอบ",
        "lesson_id → course_lessons.id : ระบุบทเรียนที่ข้อสอบนี้เกี่ยวข้อง (Many-to-One, Nullable)",
        "tutor_id → tutors.id : ระบุครูผู้สร้างและดูแลข้อสอบ (Many-to-One)"
    ]
})

# ──────────────────────────────────────────────
# 19. exam_questions
# ──────────────────────────────────────────────
TABLES.append({
    "name": "exam_questions",
    "thai": "19. คำถามในข้อสอบ",
    "rows": [
        [1,  "id",             "BIGINT",    "8",   "รหัสคำถาม (Auto Increment)",                                                           "PK"],
        [2,  "exam_id",        "BIGINT",    "8",   "อ้างอิงข้อสอบที่คำถามนี้สังกัด",                                                    "FK → exams.id"],
        [3,  "question_text",  "TEXT",      "-",   "ข้อความของคำถาม",                                                                     ""],
        [4,  "question_type",  "VARCHAR",   "30",  "ประเภท: MULTIPLE_CHOICE, SHORT_ANSWER, PARAGRAPH, CHECKBOX",                         ""],
        [5,  "explanation",    "TEXT",      "-",   "คำอธิบายเฉลยสำหรับแสดงหลังส่งข้อสอบ",                                              ""],
        [6,  "score",          "DOUBLE",    "-",   "คะแนนสำหรับคำถามข้อนี้",                                                            ""],
        [7,  "is_required",    "BOOLEAN",   "-",   "บังคับต้องตอบหรือไม่ (ค่าเริ่มต้น: true)",                                          ""],
        [8,  "question_order", "INT",       "4",   "ลำดับข้อคำถาม",                                                                      ""],
        [9,  "created_at",     "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                                   ""],
        [10, "updated_at",     "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                                          ""],
    ],
    "fk_notes": [
        "exam_id → exams.id : คำถามสังกัดอยู่ในข้อสอบ (Many-to-One) ข้อสอบหนึ่งมีได้หลายคำถาม เมื่อลบข้อสอบจะลบคำถามทั้งหมดด้วย (Cascade)"
    ]
})

# ──────────────────────────────────────────────
# 20. exam_question_options
# ──────────────────────────────────────────────
TABLES.append({
    "name": "exam_question_options",
    "thai": "20. ตัวเลือกคำถามในข้อสอบ",
    "rows": [
        [1, "id",           "BIGINT",    "8",   "รหัสตัวเลือก (Auto Increment)",                              "PK"],
        [2, "question_id",  "BIGINT",    "8",   "อ้างอิงคำถามที่ตัวเลือกนี้สังกัด",                        "FK → exam_questions.id"],
        [3, "option_text",  "TEXT",      "-",   "ข้อความของตัวเลือก",                                        ""],
        [4, "is_correct",   "BOOLEAN",   "-",   "ระบุว่าตัวเลือกนี้เป็นคำตอบที่ถูกต้องหรือไม่",            ""],
        [5, "option_order", "INT",       "4",   "ลำดับตัวเลือก",                                             ""],
        [6, "created_at",   "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                          ""],
        [7, "updated_at",   "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                 ""],
    ],
    "fk_notes": [
        "question_id → exam_questions.id : ตัวเลือกสังกัดอยู่ในคำถาม (Many-to-One) คำถามหนึ่งมีได้หลายตัวเลือก เมื่อลบคำถามจะลบตัวเลือกทั้งหมดด้วย (Cascade)"
    ]
})

# ──────────────────────────────────────────────
# 21. exam_submissions
# ──────────────────────────────────────────────
TABLES.append({
    "name": "exam_submissions",
    "thai": "21. การส่งข้อสอบ",
    "rows": [
        [1,  "id",               "BIGINT",    "8",   "รหัสการส่งข้อสอบ (Auto Increment)",                                 "PK"],
        [2,  "submission_code",  "VARCHAR",   "50",  "รหัสการส่งข้อสอบ (ไม่ซ้ำกัน)",                                    "UNIQUE"],
        [3,  "exam_id",          "BIGINT",    "8",   "อ้างอิงข้อสอบที่ส่ง",                                             "FK → exams.id"],
        [4,  "student_id",       "BIGINT",    "8",   "อ้างอิงนักเรียนผู้ส่งข้อสอบ",                                    "FK → students.id"],
        [5,  "enrollment_id",    "BIGINT",    "8",   "อ้างอิงการลงทะเบียน (ตรวจสอบสิทธิ์ทำข้อสอบ)",                  "FK → enrollments.id"],
        [6,  "attempt_number",   "INT",       "4",   "ครั้งที่ทำข้อสอบ (เริ่มต้นที่ 1)",                               ""],
        [7,  "started_at",       "TIMESTAMP", "-",   "วันเวลาเริ่มทำข้อสอบ (Auto Set)",                                 ""],
        [8,  "submitted_at",     "TIMESTAMP", "-",   "วันเวลาส่งข้อสอบ",                                                ""],
        [9,  "total_score",      "DOUBLE",    "-",   "คะแนนเต็มของข้อสอบ",                                              ""],
        [10, "obtained_score",   "DOUBLE",    "-",   "คะแนนที่ได้รับ (ค่าเริ่มต้น: 0)",                                ""],
        [11, "correct_count",    "INT",       "4",   "จำนวนข้อที่ตอบถูก",                                               ""],
        [12, "wrong_count",      "INT",       "4",   "จำนวนข้อที่ตอบผิด",                                               ""],
        [13, "unanswered_count", "INT",       "4",   "จำนวนข้อที่ไม่ตอบ",                                               ""],
        [14, "is_passed",        "BOOLEAN",   "-",   "ผ่านเกณฑ์หรือไม่ (null = ยังไม่ส่ง)",                           ""],
        [15, "status",           "VARCHAR",   "20",  "สถานะ: IN_PROGRESS, SUBMITTED, GRADED, TIMED_OUT",               ""],
        [16, "created_at",       "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                ""],
        [17, "updated_at",       "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                       ""],
    ],
    "fk_notes": [
        "exam_id → exams.id : การส่งผูกกับข้อสอบ (Many-to-One) ข้อสอบหนึ่งมีการส่งได้หลายครั้ง (หากอนุญาต)",
        "student_id → students.id : ระบุนักเรียนผู้ส่งข้อสอบ (Many-to-One)",
        "enrollment_id → enrollments.id : ตรวจสอบว่านักเรียนลงทะเบียนคอร์สนี้และมีสิทธิ์ทำข้อสอบ (Many-to-One)"
    ]
})

# ──────────────────────────────────────────────
# 22. exam_answers
# ──────────────────────────────────────────────
TABLES.append({
    "name": "exam_answers",
    "thai": "22. คำตอบข้อสอบ",
    "rows": [
        [1,  "id",                  "BIGINT",    "8",   "รหัสคำตอบ (Auto Increment)",                                                        "PK"],
        [2,  "submission_id",       "BIGINT",    "8",   "อ้างอิงการส่งข้อสอบที่คำตอบนี้สังกัด",                                           "FK → exam_submissions.id"],
        [3,  "question_id",         "BIGINT",    "8",   "อ้างอิงคำถามที่ตอบ",                                                              "FK → exam_questions.id"],
        [4,  "selected_option_id",  "BIGINT",    "8",   "อ้างอิงตัวเลือกที่เลือก (สำหรับ MULTIPLE_CHOICE, Nullable)",                     "FK → exam_question_options.id"],
        [5,  "student_answer_text", "TEXT",      "-",   "คำตอบข้อความ (สำหรับ SHORT_ANSWER, PARAGRAPH หรือ CHECKBOX)",                    ""],
        [6,  "is_correct",          "BOOLEAN",   "-",   "ถูกต้องหรือไม่ (null = รอตรวจด้วยมือ)",                                          ""],
        [7,  "score_awarded",       "DOUBLE",    "-",   "คะแนนที่ได้รับสำหรับข้อนี้ (ค่าเริ่มต้น: 0)",                                   ""],
        [8,  "answered_at",         "TIMESTAMP", "-",   "วันเวลาที่ตอบคำถาม (Auto Set)",                                                   ""],
        [9,  "created_at",          "TIMESTAMP", "-",   "วันเวลาที่สร้างข้อมูล (Auto Set)",                                                ""],
        [10, "updated_at",          "TIMESTAMP", "-",   "วันเวลาที่แก้ไขข้อมูลล่าสุด (Auto Update)",                                       ""],
    ],
    "fk_notes": [
        "submission_id → exam_submissions.id : คำตอบสังกัดอยู่ในการส่งข้อสอบ (Many-to-One) เมื่อลบการส่งจะลบคำตอบทั้งหมดด้วย (Cascade)",
        "question_id → exam_questions.id : อ้างอิงคำถามที่นักเรียนตอบ (Many-to-One)",
        "selected_option_id → exam_question_options.id : อ้างอิงตัวเลือกที่นักเรียนเลือก (Many-to-One, Nullable) — Nullable เพราะ SHORT_ANSWER และ PARAGRAPH ไม่มีตัวเลือก"
    ]
})

# ──────────────────────────────────────────────
# 23. exam_score_audit_logs
# ──────────────────────────────────────────────
TABLES.append({
    "name": "exam_score_audit_logs",
    "thai": "23. บันทึกการแก้ไขคะแนนข้อสอบ",
    "rows": [
        [1, "id",            "BIGINT",    "8",   "รหัส Log (Auto Increment)",                              "PK"],
        [2, "submission_id", "BIGINT",    "8",   "อ้างอิงการส่งข้อสอบที่ถูกแก้ไขคะแนน",               "FK → exam_submissions.id"],
        [3, "question_id",   "BIGINT",    "8",   "รหัสคำถามที่ถูกแก้ไขคะแนน (ไม่มี FK constraint)",   ""],
        [4, "old_score",     "DOUBLE",    "-",   "คะแนนเดิมก่อนแก้ไข",                                  ""],
        [5, "new_score",     "DOUBLE",    "-",   "คะแนนใหม่หลังแก้ไข",                                  ""],
        [6, "changed_by",    "VARCHAR",   "255", "ผู้ที่ดำเนินการแก้ไขคะแนน",                           ""],
        [7, "reason",        "TEXT",      "-",   "เหตุผลในการแก้ไขคะแนน",                               ""],
        [8, "changed_at",    "TIMESTAMP", "-",   "วันเวลาที่แก้ไขคะแนน (Auto Set)",                     ""],
    ],
    "fk_notes": [
        "submission_id → exam_submissions.id : Log ผูกกับการส่งข้อสอบ (Many-to-One) ทุกครั้งที่แก้ไขคะแนนจะสร้าง Log ใหม่เพื่อความโปร่งใสและตรวจสอบย้อนหลังได้"
    ]
})


# ============================================================
# Build Document
# ============================================================
import docx
from docx.enum.text import WD_BREAK

doc = Document()

# ตั้งค่าหน้าแนวนอน (Landscape A4)
section = doc.sections[0]
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin  = Cm(2.0)
section.right_margin = Cm(1.5)
section.top_margin   = Cm(2.0)
section.bottom_margin = Cm(1.5)

# ===================== หน้าปก =====================
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_title.paragraph_format.space_before = Pt(120)
r = cover_title.add_run("Data Dictionary")
r.font.name = FONT_NAME
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(TITLE_COLOR)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("ระบบสถาบันกวดวิชา (Tutor School System)")
r2.font.name = FONT_NAME
r2.font.size = Pt(24)
r2.font.bold = True
r2.font.color.rgb = RGBColor.from_string("2E4057")

doc.add_paragraph()
info_lines = [
    "Backend: Spring Boot 3.5 · Java 21 · PostgreSQL",
    "จำนวนตาราง: 23 ตาราง",
    f"วันที่จัดทำ: 25 มิถุนายน 2569",
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p.add_run(line)
    r3.font.name = FONT_NAME
    r3.font.size = Pt(16)
    r3.font.color.rgb = RGBColor.from_string("555555")

# Page break หลังหน้าปก
pb = doc.add_paragraph()
pb.add_run().add_break(WD_BREAK.PAGE)

# ===================== หมายเหตุ Key =====================
note_title = doc.add_paragraph()
nt_run = note_title.add_run("สัญลักษณ์ที่ใช้ใน Data Dictionary")
nt_run.font.name = FONT_NAME
nt_run.font.size = Pt(18)
nt_run.font.bold = True
nt_run.font.color.rgb = RGBColor.from_string(TITLE_COLOR)

legend = [
    ("PK", "Primary Key — รหัสหลักของตาราง ไม่ซ้ำกัน และไม่เป็น NULL"),
    ("FK", "Foreign Key — อ้างอิงข้อมูลในตารางอื่น"),
    ("UNIQUE", "ค่าในคอลัมน์นี้ไม่ซ้ำกันในทุกแถว"),
    ("-", "ไม่มีการกำหนดขนาดเฉพาะ (TEXT, BOOLEAN, TIMESTAMP, DATE ฯลฯ)"),
    ("Nullable", "ยอมรับค่า NULL (ไม่บังคับกรอก)"),
    ("Auto Set", "ระบบกำหนดค่าให้อัตโนมัติ ไม่ต้องกรอก"),
]
for key, desc in legend:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    bold_run = p.add_run(f"• {key}  ")
    bold_run.font.name = FONT_NAME
    bold_run.font.size = Pt(14)
    bold_run.font.bold = True
    desc_run = p.add_run(desc)
    desc_run.font.name = FONT_NAME
    desc_run.font.size = Pt(14)

pb2 = doc.add_paragraph()
pb2.add_run().add_break(WD_BREAK.PAGE)

# ===================== ตารางทั้งหมด =====================
for t in TABLES:
    add_table(doc, t["name"], t["thai"], t["rows"], t.get("fk_notes"))

# บันทึกไฟล์
output_path = r"C:\Users\kakmi\tutor-school-system\Data_Dictionary_TutorSchool.docx"
doc.save(output_path)
print(f"✅ บันทึกไฟล์เรียบร้อย: {output_path}")
